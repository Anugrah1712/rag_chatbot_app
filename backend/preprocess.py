#preprocess.py 
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
from fastapi import UploadFile
import asyncio
from playwright.async_api import async_playwright
import os 
from dotenv import load_dotenv
from io import BytesIO

import os 
from dotenv import load_dotenv
# Load environment variables from .env file
load_dotenv()


async def preprocess_text(files: list[UploadFile], scraped_data, size, overlap):
    import time
    
    paragraphs = []
    from io import BytesIO
    # Step 1: Process each file
    for file in files:
        if file.filename.endswith(".pdf"):
            contents = await file.read()
            file_object = BytesIO(contents)
            reader = PdfReader(file_object)
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    cleaned_text = ' '.join(page_text.split())
                    paragraphs.append(cleaned_text)
        elif file.filename.endswith(".docx"):
            
            file_content = await file.read() 
            docx_stream = BytesIO(file_content) 
            docx = DocxDocument(docx_stream)
            full_text = ""
            for para in docx.paragraphs:
                if para.text.strip():
                    full_text += para.text.strip() + "\n\n"
            paragraphs.append(full_text)

        # Step 2: Use Playwright for web scraping
        # async with async_playwright() as p:
        #     browser = await p.chromium.launch(headless=True)  # <-- Await launch()
        #     context = await browser.new_context()  # <-- Await new_context()
        #     page = await context.new_page()  # <-- Await new_page()

        #     for link in links:
        #         try:
        #             await page.goto(link, timeout=60000)  # <-- Await page navigation
        #             await asyncio.sleep(3)  # Allow page to load
        #             body_text = await page.text_content("body")  # <-- Await text extraction
        #             paragraphs.extend(body_text.split("\n"))

        #         except Exception as link_error:
        #             print(f"Failed to process link {link}: {link_error}")

        #     await browser.close()  # <-- Await browser close

    if scraped_data:
        print(f"✅ Adding {len(scraped_data)} scraped web pages to preprocessing.")
        paragraphs.extend(scraped_data)

    paragraphs = [para.strip() for para in paragraphs if para.strip()]
    # print(paragraphs)

    docs = [LangchainDocument(page_content=para) for para in paragraphs]

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=size, chunk_overlap=overlap)
    text_chunks = text_splitter.split_documents(docs)
    # print(text_chunks)
    return text_chunks

def preprocess_chroma(text, embedding_model_name, persist_directory):
    from langchain_community.embeddings import SentenceTransformerEmbeddings
    from langchain.vectorstores import Chroma

    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)
    vectordb = Chroma.from_documents(documents=text, embedding=embedding_model, persist_directory=persist_directory)
    #vectordb.persist()
    
    retriever = vectordb.as_retriever()

    return vectordb, retriever

def preprocess_faiss(text, embedding_model_name):
    from langchain_community.embeddings import SentenceTransformerEmbeddings
    import numpy as np
    import faiss
    from langchain.docstore.in_memory import InMemoryDocstore
    from langchain.vectorstores import FAISS
    from langchain.docstore.document import Document

    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)
    texts = [doc.page_content for doc in text]
    embeddings = embedding_model.embed_documents(texts)
    embeddings = np.array(embeddings)
    dimension = embeddings.shape[1]

    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings)

    docstore = InMemoryDocstore({i: Document(page_content=texts[i]) for i in range(len(texts))})
    index_to_docstore_id = {i: i for i in range(len(texts))}

    vector_store = FAISS(
    index=index,
    docstore=docstore,
    index_to_docstore_id=index_to_docstore_id,
    embedding_function=embedding_model.embed_query 
)
    vector_store.save_local("faiss_index")  # ✅ Save FAISS index before returning
   
    return index, docstore, index_to_docstore_id, vector_store

def preprocess_weaviate(text, embedding_model_name):
    from langchain_community.embeddings import SentenceTransformerEmbeddings
    import os
    import weaviate
    from weaviate.auth import AuthApiKey
    from langchain_weaviate.vectorstores import WeaviateVectorStore

    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)

    weaviate_url = os.getenv("WEAVIATE_URL")
    weaviate_api_key = os.getenv("WEAVIATE_API_KEY")


    client = weaviate.connect_to_weaviate_cloud(
        cluster_url=weaviate_url,
        auth_credentials=AuthApiKey(weaviate_api_key),
    )

    vs = WeaviateVectorStore.from_documents(
        documents=text,
        embedding=embedding_model,
        client=client
    )
   
    return vs


def preprocess_pinecone(text, embedding_model_name):
    import numpy as np
    import time
    from langchain_community.embeddings import SentenceTransformerEmbeddings
    from pinecone import Pinecone, ServerlessSpec
    import uuid

    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)
    texts = [doc.page_content for doc in text]
    embeddings = embedding_model.embed_documents(texts)
    embeddings = np.array(embeddings).tolist()

    # Initialize Pinecone
    pinecone = Pinecone(
        api_key=os.getenv("PINECONE_API_KEY"),
        environment="us-east-1"
    )

    index_name = "test5"
    indexes = pinecone.list_indexes().names()

    # Create index if it does not exist
    if index_name not in indexes:
        pinecone.create_index(
            name=index_name,
            dimension=len(embeddings[0]),
            metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1"),
            deletion_protection="disabled"
        )
        time.sleep(5)  # Wait for the index to be created

    # Connect to existing Pinecone index
    pinecone_index = pinecone.Index(index_name)

    # Upload embeddings in batches
    upsert_data = [(str(uuid.uuid4()), embeddings[i], {"text": texts[i]}) for i in range(len(texts))]
    batch_size = 100
    for i in range(0, len(upsert_data), batch_size):
        pinecone_index.upsert(vectors=upsert_data[i:i + batch_size])
        time.sleep(0.5)

    return index_name  # ✅ Return only index_name, not the index object

def preprocess_qdrant(text, embedding_model_name):
    from qdrant_client import QdrantClient, models
    from langchain_community.embeddings import SentenceTransformerEmbeddings
    import numpy as np
    import os

    # ✅ Load Qdrant credentials
    qdrant_url = os.getenv("QDRANT_URL")  # Default to local Qdrant
    qdrant_api_key = os.getenv("QDRANT_API_KEY", None)
    collection_name = os.getenv("QDRANT_COLLECTION_NAME")

    # ✅ Initialize Qdrant Client
    client = QdrantClient(url=qdrant_url, api_key=qdrant_api_key , timeout=120)

    # ✅ Define the embedding model
    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)

    # ✅ Prepare text embeddings
    texts = [doc.page_content for doc in text]
    embeddings = embedding_model.embed_documents(texts)
    embeddings = np.array(embeddings)

    # ✅ Create a Qdrant collection if not exists
    client.recreate_collection(
        collection_name=collection_name,
        vectors_config=models.VectorParams(size=embeddings.shape[1], distance=models.Distance.COSINE)
    )

    # ✅ Insert documents in batches
    batch_size = 250
    for i in range(0, len(texts), batch_size):
        batch_texts = texts[i: i + batch_size]
        batch_embs = embeddings[i: i + batch_size]

        client.upsert(
            collection_name=collection_name,
            points=[
                models.PointStruct(
                    id=i + j,  # Assigning unique ID
                    vector=batch_embs[j].tolist(),
                    payload={"page_content": batch_texts[j]}
                )
                for j in range(len(batch_texts))
            ]
        )

    print(f"✅ Qdrant: {len(texts)} documents stored in `{collection_name}` collection.")

    return client  # ✅ Return Qdrant Client


from langchain_community.embeddings import SentenceTransformerEmbeddings
from fastapi import UploadFile
import sys
import sqlite3

sys.modules["sqlite3"] = sqlite3

embedding_model_global = None

async def preprocess_vectordbs(files: list[UploadFile], scraped_data, embedding_model_name, size, overlap):
    global embedding_model_global

    text = await preprocess_text(files, scraped_data, size, overlap)
    persist_directory = 'db'

    embedding_model_global = SentenceTransformerEmbeddings(model_name=embedding_model_name)

    print("Processing Chroma DB...")
    vectordb, retriever = preprocess_chroma(text, embedding_model_name, persist_directory)
    
    print("Processing FAISS...")
    index, docstore, index_to_docstore_id, vector_store = preprocess_faiss(text, embedding_model_name)

    print("Processing Pinecone...")
    pinecone_index_name = preprocess_pinecone(text, embedding_model_name)

    print("Processing Weaviate...")
    vs = preprocess_weaviate(text, embedding_model_name)  

    print("Processing Qdrant...")
    qdrant_client = preprocess_qdrant(text, embedding_model_name)

    return index, docstore, index_to_docstore_id, vector_store, retriever, embedding_model_global, pinecone_index_name, vs, qdrant_client  

